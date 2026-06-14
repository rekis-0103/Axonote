
from worker.pipeline.keywords import extract_keywords
from worker.pipeline.preprocess import split_sentences
from worker.pipeline.questions import generate_questions
from worker.pipeline.summary import textrank_summary


def test_split_sentences() -> None:
    text = "Machine learning is useful. It helps analyze text. Short."
    sentences = split_sentences(text)
    assert len(sentences) == 2
    assert sentences[0].startswith("Machine learning")


def test_textrank_summary_returns_ordered_subset() -> None:
    sentences = [
        "Neural networks learn patterns from data.",
        "Deep learning uses many layers.",
        "Training requires labeled examples.",
        "Evaluation measures model quality.",
    ]
    summary = textrank_summary(sentences, max_sentences=2)
    assert len(summary) == 2
    assert all(sentence in sentences for sentence in summary)


def test_generate_questions_builds_mcq() -> None:
    sentences = [
        "Photosynthesis converts sunlight into chemical energy in plants.",
        "Chlorophyll captures light for the photosynthesis process.",
    ]
    keywords = extract_keywords(sentences)
    questions = generate_questions(sentences, keywords, max_questions=2)
    assert questions
    first = questions[0]
    assert first["type"] == "mcq"
    assert len(first["options"]) == 4
    assert 0 <= int(first["correct_index"]) < 4


def test_run_analysis_on_sample_text(tmp_path) -> None:
    from worker.pipeline.run import run_analysis

    sample = tmp_path / "notes.pdf"
    sample.write_bytes(b"%PDF-1.4\nPhotosynthesis converts light. Plants use chlorophyll.\n")

    # pypdf may not extract well from minimal pdf - use docx instead
    docx_path = tmp_path / "notes.docx"
    from docx import Document

    document = Document()
    document.add_paragraph("Machine learning models learn from data.")
    document.add_paragraph("Supervised learning uses labeled training examples.")
    document.save(str(docx_path))

    result = run_analysis(docx_path, "docx")
    assert result["content"]
    assert isinstance(result["keywords"], list)
    assert result["questions"]
